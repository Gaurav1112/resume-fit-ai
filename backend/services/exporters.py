"""Export a TailoredResume to genuinely ATS-friendly DOCX, PDF and TXT.

Both binary exporters render from the same structured model as the plain-text
renderer, so what the validators checked is what the file contains.

DOCX: real Word paragraphs with real styles. No tables, no text boxes, no
headers/footers, no images, no columns — everything a parser needs is in the
document body as linear runs.

PDF: reportlab Platypus flowables, producing a real text layer. Not an image, not
a screenshot embedded in a page — you can select and copy the text, which means
an ATS can extract it.
"""

from __future__ import annotations

import io
import re

from ..models.schemas import TailoredResume
from .render import contact_line, to_plain_text

SECTION_ORDER = [
    "summary", "skills", "experience", "projects", "achievements",
    "education", "certifications",
]


def _ordered(resume: TailoredResume):
    order = {kind: i for i, kind in enumerate(SECTION_ORDER)}
    return sorted(resume.sections, key=lambda s: order.get(s.kind, 99))


def safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("_")
    return cleaned or "resume"


# --------------------------------------------------------------------------- #
# TXT
# --------------------------------------------------------------------------- #
def to_txt(resume: TailoredResume) -> bytes:
    return to_plain_text(resume).encode("utf-8")


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def to_docx(resume: TailoredResume) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Calibri is a Word core font and universally parseable. Set it on the
    # Normal style so every paragraph inherits it.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.06

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(46)

    def para(text: str, *, size=10.5, bold=False, italic=False, space_before=0,
             space_after=2, align=None, color=None):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def heading(text: str):
        p = para(text.upper(), size=11, bold=True, space_before=10, space_after=3)
        # A bottom border is a paragraph property, not a table — parser-safe.
        pPr = p._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        borders.append(bottom)
        pPr.append(borders)
        return p

    def bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Pt(14)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"

    # --- Header block (in the body, never in a Word header) ---
    if resume.contact.name:
        para(resume.contact.name, size=18, bold=True, space_after=1,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    if resume.headline:
        para(resume.headline, size=11.5, space_after=2,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    line = contact_line(resume.contact)
    if line:
        para(line, size=9.5, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(0x44, 0x44, 0x44))

    for section in _ordered(resume):
        if section.kind == "summary":
            body = [p for p in section.paragraphs if p.strip()]
            if not body:
                continue
            heading(section.heading)
            for text in body:
                para(text)

        elif section.kind == "skills":
            groups = {g: s for g, s in section.skill_groups.items() if s}
            if not groups:
                continue
            heading(section.heading)
            for group, skills in groups.items():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                label = p.add_run(f"{group}: ")
                label.bold = True
                label.font.size = Pt(10.5)
                value = p.add_run(", ".join(skills))
                value.font.size = Pt(10.5)

        elif section.kind == "experience":
            if not section.roles:
                continue
            heading(section.heading)
            for role in section.roles:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(0)
                title = p.add_run(f"{role.title}")
                title.bold = True
                title.font.size = Pt(11)
                company = p.add_run(f", {role.company}")
                company.font.size = Pt(11)
                meta = " | ".join(
                    x for x in (
                        f"{role.start_date} - {role.end_date}".strip(" -"), role.location
                    ) if x
                )
                if meta:
                    para(meta, size=9.5, italic=True, space_after=2,
                         color=(0x55, 0x55, 0x55))
                for b in role.bullets:
                    if b.text.strip():
                        bullet(b.text)

        elif section.kind == "education":
            if not section.education:
                continue
            heading(section.heading)
            for edu in section.education:
                head = ", ".join(x for x in (edu.degree, edu.field_of_study) if x)
                para(f"{head} — {edu.institution}" if head else edu.institution, bold=True)
                meta = " | ".join(
                    x for x in (
                        f"{edu.start_date} - {edu.end_date}".strip(" -"), edu.location
                    ) if x
                )
                if meta:
                    para(meta, size=9.5, italic=True, color=(0x55, 0x55, 0x55))
                for detail in edu.details:
                    bullet(detail)

        elif section.kind == "certifications":
            if not section.certifications:
                continue
            heading(section.heading)
            for cert in section.certifications:
                bits = [b for b in (cert.name, cert.issuer, cert.date) if b]
                bullet(" — ".join(bits))

        else:
            items = [b for b in section.bullets if b.text.strip()]
            if not items and not section.paragraphs:
                continue
            heading(section.heading)
            for text in section.paragraphs:
                para(text)
            for b in items:
                bullet(b.text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def to_pdf(resume: TailoredResume) -> bytes:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer,
    )
    from xml.sax.saxutils import escape

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=0.62 * inch,
        rightMargin=0.62 * inch,
        topMargin=0.52 * inch,
        bottomMargin=0.52 * inch,
        title=f"{resume.contact.name or 'Resume'}",
        author=resume.contact.name or "",
    )

    base = getSampleStyleSheet()
    styles = {
        "name": ParagraphStyle(
            "name", parent=base["Title"], fontName="Helvetica-Bold", fontSize=18,
            leading=21, alignment=TA_CENTER, spaceAfter=1,
        ),
        "headline": ParagraphStyle(
            "headline", parent=base["Normal"], fontSize=11, leading=13,
            alignment=TA_CENTER, spaceAfter=2,
        ),
        "contact": ParagraphStyle(
            "contact", parent=base["Normal"], fontSize=8.8, leading=11,
            alignment=TA_CENTER, textColor="#444444", spaceAfter=6,
        ),
        "heading": ParagraphStyle(
            "heading", parent=base["Heading2"], fontName="Helvetica-Bold",
            fontSize=10.5, leading=12.5, spaceBefore=9, spaceAfter=1,
            textColor="#111111",
        ),
        "body": ParagraphStyle(
            "body", parent=base["Normal"], fontSize=9.6, leading=12.4, spaceAfter=2,
        ),
        "role": ParagraphStyle(
            "role", parent=base["Normal"], fontSize=10.4, leading=12.6,
            spaceBefore=5, spaceAfter=0,
        ),
        "meta": ParagraphStyle(
            "meta", parent=base["Normal"], fontSize=8.8, leading=10.6,
            textColor="#555555", spaceAfter=2,
        ),
        "bullet": ParagraphStyle(
            "bullet", parent=base["Normal"], fontSize=9.6, leading=12.2, spaceAfter=1.5,
        ),
    }

    def p(text: str, style: str):
        return Paragraph(escape(text), styles[style])

    story: list = []
    if resume.contact.name:
        story.append(p(resume.contact.name, "name"))
    if resume.headline:
        story.append(p(resume.headline, "headline"))
    line = contact_line(resume.contact)
    if line:
        story.append(p(line, "contact"))

    def section_heading(text: str):
        story.append(p(text.upper(), "heading"))
        story.append(HRFlowable(width="100%", thickness=0.5, color="#999999",
                                spaceBefore=1, spaceAfter=4))

    def bullets(texts: list[str]):
        items = [ListItem(p(t, "bullet"), leftIndent=10) for t in texts if t.strip()]
        if items:
            story.append(
                ListFlowable(items, bulletType="bullet", start="•", leftIndent=12,
                             bulletFontSize=7, spaceBefore=0, spaceAfter=2)
            )

    for section in _ordered(resume):
        if section.kind == "summary":
            body = [t for t in section.paragraphs if t.strip()]
            if not body:
                continue
            section_heading(section.heading)
            for text in body:
                story.append(p(text, "body"))

        elif section.kind == "skills":
            groups = {g: s for g, s in section.skill_groups.items() if s}
            if not groups:
                continue
            section_heading(section.heading)
            for group, skills in groups.items():
                story.append(
                    Paragraph(
                        f"<b>{escape(group)}:</b> {escape(', '.join(skills))}",
                        styles["body"],
                    )
                )

        elif section.kind == "experience":
            if not section.roles:
                continue
            section_heading(section.heading)
            for role in section.roles:
                story.append(
                    Paragraph(
                        f"<b>{escape(role.title)}</b>, {escape(role.company)}",
                        styles["role"],
                    )
                )
                meta = " | ".join(
                    x for x in (
                        f"{role.start_date} - {role.end_date}".strip(" -"), role.location
                    ) if x
                )
                if meta:
                    story.append(p(meta, "meta"))
                bullets([b.text for b in role.bullets])

        elif section.kind == "education":
            if not section.education:
                continue
            section_heading(section.heading)
            for edu in section.education:
                head = ", ".join(x for x in (edu.degree, edu.field_of_study) if x)
                story.append(
                    Paragraph(
                        f"<b>{escape(head or edu.institution)}</b>"
                        + (f" — {escape(edu.institution)}" if head else ""),
                        styles["body"],
                    )
                )
                meta = " | ".join(
                    x for x in (
                        f"{edu.start_date} - {edu.end_date}".strip(" -"), edu.location
                    ) if x
                )
                if meta:
                    story.append(p(meta, "meta"))
                bullets(edu.details)

        elif section.kind == "certifications":
            if not section.certifications:
                continue
            section_heading(section.heading)
            bullets([
                " — ".join(x for x in (c.name, c.issuer, c.date) if x)
                for c in section.certifications
            ])

        else:
            items = [b.text for b in section.bullets if b.text.strip()]
            if not items and not section.paragraphs:
                continue
            section_heading(section.heading)
            for text in section.paragraphs:
                story.append(p(text, "body"))
            bullets(items)

    story.append(Spacer(1, 2))
    doc.build(story)
    return buffer.getvalue()


EXPORTERS = {
    "txt": (to_txt, "text/plain; charset=utf-8", "txt"),
    "docx": (
        to_docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "pdf": (to_pdf, "application/pdf", "pdf"),
}
