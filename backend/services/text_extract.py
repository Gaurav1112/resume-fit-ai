"""Extract normalised plain text from PDF / DOCX / TXT / Markdown uploads."""

from __future__ import annotations

import io
import re
import unicodedata

SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".markdown", ".rtf", ".html", ".htm"}


class ExtractionError(ValueError):
    pass


def _clean(text: str) -> str:
    """Normalise whitespace, unicode punctuation and bullet glyphs."""
    text = unicodedata.normalize("NFKC", text)
    replacements = {
        "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "-", "\u00a0": " ", "\ufeff": "",
        "\u2022": "-", "\u25cf": "-", "\u25aa": "-", "\u00b7": "-", "\u2023": "-",
        "\uf0b7": "-", "\uf0a7": "-",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [ln.rstrip() for ln in text.splitlines()]
    return "\n".join(lines).strip()


def _from_pdf(data: bytes) -> str:
    """pdfplumber preserves column/table layout better; pypdf is the fallback."""
    text = ""
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        text = "\n\n".join(pages).strip()
    except Exception:  # noqa: BLE001 - fall through to pypdf
        text = ""

    if len(text) < 120:
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()
        except Exception as exc:  # noqa: BLE001
            raise ExtractionError(f"Could not read the PDF: {exc}") from exc

    if len(text) < 60:
        raise ExtractionError(
            "This PDF contains almost no extractable text — it is probably a scan or "
            "an image export. Upload a text-based PDF, a DOCX, or paste the text. "
            "(A PDF an ATS cannot read is itself a serious ATS problem.)"
        )
    return text


def _from_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError("python-docx is not installed.") from exc

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs]
    # Tables are ATS-hostile but common — read them so we don't lose content.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


_BLOCK_TAGS = (
    "p", "div", "li", "tr", "br", "h1", "h2", "h3", "h4", "h5", "h6",
    "section", "article", "header", "footer", "ul", "ol", "table",
)


def _from_html(data: bytes) -> str:
    """Extract text from an HTML resume.

    Many people keep their master resume as an HTML file they render to PDF, so
    this is a first-class input. Block-level tags become newlines and list items
    become bullets, which is what the downstream section parser expects; script,
    style and head content is dropped entirely.
    """
    html = data.decode("utf-8", errors="replace")

    html = re.sub(r"(?is)<(script|style|head|noscript)\b.*?</\1\s*>", " ", html)
    html = re.sub(r"(?s)<!--.*?-->", " ", html)

    # List items become bullets before the generic tag strip, so the resume
    # parser can still tell a bullet from a paragraph.
    html = re.sub(r"(?i)<li\b[^>]*>", "\n- ", html)
    for tag in _BLOCK_TAGS:
        html = re.sub(rf"(?i)</?{tag}\b[^>]*>", "\n", html)
    html = re.sub(r"(?i)<t[dh]\b[^>]*>", " | ", html)
    html = re.sub(r"<[^>]+>", "", html)

    import html as html_module

    text = html_module.unescape(html)
    # Collapse the runs of blank lines that tag stripping leaves behind, but keep
    # single blank lines — the projects parser uses them as separators.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    lines = [ln.strip() for ln in text.splitlines()]
    out: list[str] = []
    for line in lines:
        if line or (out and out[-1]):
            out.append(line)
    return "\n".join(out)


def _from_rtf(data: bytes) -> str:
    raw = data.decode("utf-8", errors="ignore")
    raw = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: chr(int(m.group(1), 16)), raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
    return raw.replace("{", " ").replace("}", " ")


def extract(filename: str, data: bytes) -> str:
    """Dispatch on extension; sniff the PDF magic number as a safety net."""
    name = (filename or "").lower()
    suffix = "." + name.rsplit(".", 1)[-1] if "." in name else ""

    if data[:5] == b"%PDF-":
        return _clean(_from_pdf(data))
    if data[:2] == b"PK" and suffix == ".docx":
        return _clean(_from_docx(data))

    if suffix == ".pdf":
        return _clean(_from_pdf(data))
    if suffix == ".docx":
        return _clean(_from_docx(data))
    if suffix in {".html", ".htm"}:
        return _clean(_from_html(data))
    if suffix == ".rtf":
        return _clean(_from_rtf(data))
    if data[:200].lstrip()[:9].lower().startswith((b"<!doctype", b"<html")):
        return _clean(_from_html(data))
    if suffix in {".txt", ".md", ".markdown", ""}:
        return _clean(data.decode("utf-8", errors="replace"))
    if suffix == ".doc":
        raise ExtractionError(
            "Legacy .doc is not supported. Save as .docx or PDF and re-upload."
        )
    raise ExtractionError(
        f"Unsupported file type '{suffix}'. Supported: PDF, DOCX, TXT, Markdown, RTF."
    )
